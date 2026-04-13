#!/usr/bin/env python3
"""Unit tests for docx_helper.py"""

import json
import os
import sys
import tempfile
import unittest

# Add scripts dir to path
sys.path.insert(0, os.path.dirname(__file__))

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx_helper

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
HELPER = os.path.join(SCRIPT_DIR, "docx_helper.py")


def run_helper(*args):
    """Run the helper script and return (exit_code, stdout, stderr)."""
    import subprocess
    result = subprocess.run(
        ["python3", HELPER] + list(args),
        capture_output=True, text=True
    )
    return result.returncode, result.stdout, result.stderr


def create_simple_docx(path, paragraphs):
    """Create a simple DOCX with the given paragraph texts."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


def create_table_docx(path):
    """Create a DOCX with a table containing paragraphs."""
    doc = Document()
    doc.add_paragraph("Before table")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cell A"
    table.cell(0, 1).text = "Cell B"
    table.cell(1, 0).text = "Cell C"
    table.cell(1, 1).text = "Cell D"
    doc.add_paragraph("After table")
    doc.save(path)


def create_multi_run_docx(path):
    """Create a DOCX with a paragraph that has multiple runs with different formatting."""
    doc = Document()
    p = doc.add_paragraph()
    run1 = p.add_run("Bold part")
    run1.bold = True
    run2 = p.add_run(" and normal part")
    doc.save(path)


class TestExtract(unittest.TestCase):
    def test_simple_extract(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            path = f.name
        try:
            create_simple_docx(path, ["Hello", "World", ""])
            code, stdout, _ = run_helper("extract", path)
            self.assertEqual(code, 0)
            result = json.loads(stdout)
            self.assertEqual(len(result), 3)
            self.assertEqual(result[0]["text"], "Hello")
            self.assertEqual(result[1]["text"], "World")
            self.assertEqual(result[2]["text"], "")
        finally:
            os.unlink(path)

    def test_table_extract(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            path = f.name
        try:
            create_table_docx(path)
            code, stdout, _ = run_helper("extract", path)
            self.assertEqual(code, 0)
            result = json.loads(stdout)
            # "Before table" (0) + 4 cells (1-4) + "After table" (5)
            self.assertEqual(len(result), 6)
            self.assertEqual(result[0]["text"], "Before table")
            self.assertEqual(result[1]["text"], "Cell A")
            self.assertEqual(result[5]["text"], "After table")
        finally:
            os.unlink(path)

    def test_multi_run_extract(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            path = f.name
        try:
            create_multi_run_docx(path)
            code, stdout, _ = run_helper("extract", path)
            self.assertEqual(code, 0)
            result = json.loads(stdout)
            self.assertEqual(len(result), 1)
            self.assertEqual(result[0]["text"], "Bold part and normal part")
        finally:
            os.unlink(path)


class TestReplace(unittest.TestCase):
    def test_single_run_replace(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            in_path = f.name
        out_path = in_path + '.out'
        trans_path = in_path + '.json'
        try:
            create_simple_docx(in_path, ["Hello", "World"])
            with open(trans_path, 'w') as f:
                json.dump({"0": "你好", "1": "世界"}, f)
            code, _, stderr = run_helper("replace", in_path, out_path, trans_path)
            self.assertEqual(code, 0, f"replace failed: {stderr}")
            code, stdout, _ = run_helper("extract", out_path)
            result = json.loads(stdout)
            self.assertEqual(result[0]["text"], "你好")
            self.assertEqual(result[1]["text"], "世界")
        finally:
            for p in [in_path, out_path, trans_path]:
                if os.path.exists(p): os.unlink(p)

    def test_table_replace(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            in_path = f.name
        out_path = in_path + '.out'
        trans_path = in_path + '.json'
        try:
            create_table_docx(in_path)
            with open(trans_path, 'w') as f:
                json.dump({"0": "表格前", "2": "单元格B"}, f)
            code, _, stderr = run_helper("replace", in_path, out_path, trans_path)
            self.assertEqual(code, 0, f"replace failed: {stderr}")
            code, stdout, _ = run_helper("extract", out_path)
            result = json.loads(stdout)
            self.assertEqual(result[0]["text"], "表格前")
            self.assertEqual(result[1]["text"], "Cell A")  # not translated
            self.assertEqual(result[2]["text"], "单元格B")
            self.assertEqual(result[5]["text"], "After table")  # not translated
        finally:
            for p in [in_path, out_path, trans_path]:
                if os.path.exists(p): os.unlink(p)

    def test_multi_run_replace(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            in_path = f.name
        out_path = in_path + '.out'
        trans_path = in_path + '.json'
        try:
            create_multi_run_docx(in_path)
            with open(trans_path, 'w') as f:
                json.dump({"0": "全部翻译文本"}, f)
            code, _, stderr = run_helper("replace", in_path, out_path, trans_path)
            self.assertEqual(code, 0, f"replace failed: {stderr}")
            code, stdout, _ = run_helper("extract", out_path)
            result = json.loads(stdout)
            self.assertEqual(result[0]["text"], "全部翻译文本")
        finally:
            for p in [in_path, out_path, trans_path]:
                if os.path.exists(p): os.unlink(p)

    def test_empty_paragraph_preserved(self):
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            in_path = f.name
        out_path = in_path + '.out'
        trans_path = in_path + '.json'
        try:
            create_simple_docx(in_path, ["Hello", "", "World"])
            with open(trans_path, 'w') as f:
                json.dump({"0": "你好", "2": "世界"}, f)
            code, _, stderr = run_helper("replace", in_path, out_path, trans_path)
            self.assertEqual(code, 0, f"replace failed: {stderr}")
            code, stdout, _ = run_helper("extract", out_path)
            result = json.loads(stdout)
            self.assertEqual(len(result), 3)  # empty paragraph still there
            self.assertEqual(result[0]["text"], "你好")
            self.assertEqual(result[1]["text"], "")  # preserved
            self.assertEqual(result[2]["text"], "世界")
        finally:
            for p in [in_path, out_path, trans_path]:
                if os.path.exists(p): os.unlink(p)


if __name__ == '__main__':
    unittest.main()
